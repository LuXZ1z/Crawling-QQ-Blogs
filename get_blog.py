import csv
import time
import requests
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.common.exceptions import WebDriverException
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from io import BytesIO
import pandas as pd

qq = '123456'           # 键入你的QQ号

# 读取qq_blog_id.csv文件
df = pd.read_csv("qq_blog_link.csv")
# 获取日志ID
blog_ids = df["Blog Link"].tolist()

url = f"https://user.qzone.qq.com/{qq}/blog/"

driver = webdriver.Chrome()

cnt = 0
for blog_id in  blog_ids:
    url_ = url + str(blog_id)

    # 访问网页失败的最大尝试次数
    max_attempts = 3
    attempt = 1

    cnt += 1
    print(f"正在爬取第{cnt}篇日志{blog_id}")

    while attempt <= max_attempts:
        try:
            driver.get(url_)
            # time.sleep(5)

            iframe_element = driver.find_element(By.XPATH, "//iframe[@id='tblog']")
            driver.switch_to.frame(iframe_element)

            title = driver.find_element(By.CSS_SELECTOR, ".blog_tit_detail").text
            time = driver.find_element(By.CSS_SELECTOR, "#pubTime").text
            text = driver.find_element(By.CSS_SELECTOR, "#blogDetailDiv").text
            # imgs = driver.find_elements(By.CSS_SELECTOR, "img[alt='图片']")
            imgs = driver.find_elements(By.CSS_SELECTOR, "img")

            comments = driver.find_elements(By.CSS_SELECTOR, "#commentListDiv")

            img_list = []
            for img in imgs:
                img_list.append([img.get_attribute('data-img-idx'), img.get_attribute('data-src')])

            comment_list = []
            for comment in comments:
                comment_list.append(comment.text)

            # 将标题中的特殊字符替换为合适的字符，例如去除冒号 ":"，斜杠 "/" 和问号 "?"
            # 可以使用 title.replace() 方法进行替换，如有其他特殊字符，也可以在这里添加对应的替换规则
            title_ = title.replace(":", "").replace("/", "").replace("?", "").replace("“", '').replace("\"",'')
            time_ = time.replace(":", "").replace("/", "").replace("?", "")
            # 将内容保存到CSV文件，并使用标题作为文件名
            filename_csv = f"./csv/{title_}_{time_}.csv"
            with open(filename_csv, 'w', newline='', encoding='utf-8') as csvfile:      # 请用记事本打开，别用excel
                csvwriter = csv.writer(csvfile)

                # 写入标题、时间和文本内容
                csvwriter.writerow(['Title', 'Time', 'Text'])
                csvwriter.writerow([title, time, text])

                # 写入图片链接
                csvwriter.writerow(['Image Index', 'Image Link'])
                for img in img_list:
                    csvwriter.writerow(img)

                # 写入评论列表
                csvwriter.writerow(['Comments'])
                for comment in comment_list:
                    csvwriter.writerow([comment])

            print(f"数据已保存到文件: {filename_csv}")

            # 创建一个新的Word文档
            doc = Document()

            # 创建一个样式对象并设置字体为微软雅黑
            font_style = doc.styles['Normal'].font
            font_style.name = '微软雅黑'
            font_style.size = Pt(12)  # 设置字体大小，这里设置为12磅

            # 设置中文字体样式为微软雅黑
            style_zh = doc.styles.add_style('zh', WD_STYLE_TYPE.CHARACTER)
            style_zh.font.name = '微软雅黑'
            doc.styles['zh']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


            # 添加标题、时间和文本内容到文档
            doc.add_heading(title, level=1)
            doc.add_paragraph(time)
            doc.add_paragraph(text)

            # 添加图片链接到文档
            doc.add_heading('Image Links', level=2)
            for img in img_list:
                img_index = img[0]
                img_url = img[1]

                if img_url == None:
                    continue

                try:
                    # 下载图片
                    response = requests.get(img_url)
                    response.raise_for_status()  # 检查是否下载成功，如果返回码不是2xx，会引发HTTPError
                    image = BytesIO(response.content)

                    # 插入图片到文档中，只有下载成功的图片才进行插入操作
                    if image:
                        doc.add_paragraph(f"Image Index: {img_index}")
                        doc.add_picture(image, width=Inches(4))  # 设置图片宽度为4英寸

                    else:
                        doc.add_paragraph(f"Image Index: {img_index} (Invalid Image)")
                except requests.exceptions.RequestException as e:
                    # 下载图片出错时，设置image为None，同时在文档中标记图片为无效图片
                    image = None
                    doc.add_paragraph(f"Image Index: {img_index} (Invalid Image)")

            # 添加评论列表到文档
            doc.add_heading('Comments', level=2)
            for comment in comment_list:
                doc.add_paragraph(comment)

            # 保存文档
            filename_docx = f"./docx/{title_}_{time_}.docx"
            doc.save(filename_docx)
            print(f"数据已保存到文件: {filename_docx}")

            break

        except WebDriverException as e:
            print(f"访问网页失败，尝试次数：{attempt}/{max_attempts}")
            attempt += 1
            if attempt > max_attempts:
                print(f"达到最大尝试次数，无法访问网页{blog_id}。")
                with open('error.csv', 'a', newline='', encoding='utf-8') as csvfile:
                    csvwriter = csv.writer(csvfile)
                    csvwriter.writerow(blog_id)

                break  # 达到最大尝试次数后退出循环